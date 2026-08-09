from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
DOCX_PATH = OUT_DIR / "智慧翼企业福利商城全量技术方案与报价评估报告-雍彻科技.docx"
MD_PATH = OUT_DIR / "智慧翼企业福利商城全量技术方案与报价评估报告-雍彻科技.md"

BLUE = "1F5EFF"
NAVY = "143A8F"
ORANGE = "FF7A00"
INK = "1F2937"
MUTED = "667085"
LIGHT_BLUE = "EAF1FF"
LIGHT_GRAY = "F2F4F7"
GREEN = "18A058"
RED = "C62828"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


DOMAIN_TOPICS = {
    "DOM": ("域名、品牌与备案", [
        "主域名所有权核验", "裸域名解析", "www子域名解析", "HTTPS证书自动续期", "域名到期提醒",
        "DNS变更审计", "品牌中文名称确认", "品牌英文名称确认", "Logo与图标规范", "ICP备案主体确认",
        "公安联网备案", "隐私政策域名一致性",
    ]),
    "TEN": ("平台、集团、商城与员工四级体系", [
        "平台级租户管理", "集团级组织管理", "商城级独立配置", "员工级账户归属", "集团与商城一对多关系",
        "商城切换权限", "租户数据强隔离", "商城数据强隔离", "跨集团访问禁止", "组织架构同步",
        "商城停用与恢复", "租户数据导出与注销",
    ]),
    "IAM": ("身份、登录与权限", [
        "手机号验证码登录", "企业单点登录SSO", "微信授权登录", "账号密码策略", "多因素认证",
        "角色权限RBAC", "数据权限ABAC", "会话超时", "异地登录告警", "账号冻结解冻",
        "离职员工自动停用", "管理员操作二次确认",
    ]),
    "USR": ("员工与个人中心", [
        "员工档案", "工号唯一性", "部门岗位信息", "实名信息最小化", "收货地址管理",
        "默认地址规则", "收藏夹", "浏览历史", "消息中心", "个人发票抬头",
        "账户安全设置", "个人数据下载与注销",
    ]),
    "PC": ("PC Web商城前台", [
        "京东式高密度首页", "顶部企业状态栏", "多级商品分类", "全局搜索", "活动轮播",
        "商品瀑布流", "商品详情", "购物车", "结算台", "订单中心",
        "个人中心", "桌面端无障碍键盘操作",
    ]),
    "MOB": ("移动H5与微信小程序", [
        "美团式移动首页", "底部五栏导航", "移动分类页", "移动搜索联想", "移动商品详情",
        "移动购物车", "移动结算", "安全区域适配", "触摸目标尺寸", "微信小程序登录",
        "小程序支付", "小程序审核与发布",
    ]),
    "PRD": ("商品中心", [
        "SPU管理", "SKU管理", "规格属性", "商品图片视频", "商品详情富文本",
        "商品上下架", "库存状态", "企业专享商品", "虚拟商品类型", "服务商品类型",
        "商品审核", "商品批量导入导出",
    ]),
    "CAT": ("分类、搜索与推荐", [
        "三级类目", "类目排序", "品牌筛选", "属性筛选", "价格筛选",
        "库存筛选", "关键词纠错", "搜索热词", "搜索无结果推荐", "销量排序",
        "个性化推荐", "搜索日志分析",
    ]),
    "PRI": ("价格、促销与权益", [
        "市场价", "商城价", "福利专享价", "集团协议价", "员工分层价",
        "限时促销", "优惠券", "满减", "礼盒组合", "限购",
        "促销叠加规则", "价格变更审计",
    ]),
    "ACC": ("福利账户与预算", [
        "福利卡账户", "餐卡账户", "账户余额", "预算批次", "发放名单",
        "有效期", "使用范围", "冻结解冻", "扣减原子性", "退款原路返回",
        "账户流水", "财务对账",
    ]),
    "CHK": ("购物车与结算", [
        "购物车按商城隔离", "购物车按用户隔离", "SKU合并", "库存校验", "价格二次校验",
        "配送地址校验", "供应商拆单预览", "福利卡抵扣", "餐卡抵扣", "微信补差",
        "订单幂等提交", "防重复支付",
    ]),
    "ORD": ("订单中心", [
        "父子订单", "订单编号规则", "订单状态机", "订单快照", "订单搜索",
        "订单筛选", "取消订单", "确认收货", "订单超时关闭", "虚拟订单自动完成",
        "订单导出", "订单审计轨迹",
    ]),
    "PAY": ("支付、清结算与对账", [
        "微信支付", "支付回调验签", "支付幂等", "福利账户组合支付", "退款",
        "部分退款", "支付异常补单", "日对账", "差错处理", "供应商结算",
        "平台服务费", "财务结算单",
    ]),
    "VIR": ("虚拟卡券与核销", [
        "卡券库存", "卡券发码", "卡密加密存储", "二维码核销", "动态核销码",
        "核销门店", "使用有效期", "核销撤销", "异常卡券补发", "供应商发码回调",
        "卡券状态同步", "卡券泄露处置",
    ]),
    "FUL": ("实物履约与物流", [
        "仓库路由", "供应商发货", "物流单号", "物流轨迹", "拆包裹",
        "部分发货", "签收状态", "配送时效", "偏远地区限制", "缺货替代",
        "异常件处理", "物流订阅回调",
    ]),
    "SUP": ("供应商与分销商", [
        "供应商准入", "资质审核", "供应商合同", "分销商distributorId", "商品归属",
        "供货价", "结算周期", "服务区域", "供应商评分", "供应商停用",
        "接口密钥", "供应商数据隔离",
    ]),
    "AFS": ("售后、退款与投诉", [
        "仅退款", "退货退款", "换货", "售后原因", "凭证上传",
        "审核流程", "退货地址", "退款计算", "福利账户原路退回", "微信原路退款",
        "售后时效", "争议仲裁",
    ]),
    "INV": ("发票与税务资料", [
        "个人发票", "企业发票", "税号校验", "电子发票", "开票状态",
        "红冲", "发票下载", "发票邮件发送", "拆单开票", "服务类目税率",
        "发票抬头库", "开票审计",
    ]),
    "CS": ("客户服务制度与工单", [
        "在线客服入口", "客服电话管理", "服务时间公示", "机器人首问", "人工转接",
        "工单分类", "工单优先级", "首次响应SLA", "解决时限SLA", "升级机制",
        "满意度评价", "客服质检",
    ]),
    "OPS": ("运营与内容管理", [
        "首页装修", "活动页", "Banner管理", "公告管理", "热词管理",
        "推荐位管理", "专题页", "内容审核", "定时发布", "灰度发布",
        "运营日历", "活动复盘",
    ]),
    "ADM": ("运营管理后台", [
        "平台管理端", "集团管理端", "商城管理端", "供应商工作台", "员工管理",
        "商品审核", "订单管理", "售后管理", "账户管理", "权限管理",
        "批量任务", "导入导出中心",
    ]),
    "DAT": ("数据、报表与BI", [
        "GMV报表", "订单报表", "用户活跃", "商品转化", "搜索分析",
        "福利预算消耗", "账户余额汇总", "供应商履约", "售后率", "客服SLA",
        "数据导出", "指标口径字典",
    ]),
    "SEC": ("安全与隐私", [
        "传输加密", "敏感字段加密", "密码哈希", "密钥托管", "最小权限",
        "接口签名", "重放攻击防护", "限流", "WAF", "恶意爬虫防护",
        "隐私授权记录", "个人信息删除",
    ]),
    "CMP": ("合规与审计", [
        "隐私政策", "用户协议", "第三方SDK清单", "个人信息处理清单", "未成年人规则",
        "数据出境评估", "日志留存", "管理员审计", "资金操作审计", "供应商资质留档",
        "等保评估准备", "合规事件响应",
    ]),
    "API": ("接口与系统集成", [
        "统一API网关", "REST规范", "OpenAPI文档", "企业HR同步", "统一身份认证",
        "供应商商品接口", "库存接口", "价格接口", "订单推送", "物流回调",
        "卡券发码接口", "回调重试",
    ]),
    "DEV": ("基础设施与DevOps", [
        "开发测试生产隔离", "基础设施即代码", "自动构建", "自动测试", "自动部署",
        "版本回滚", "配置中心", "密钥管理", "数据库迁移", "对象存储",
        "CDN", "成本预算告警",
    ]),
    "TST": ("测试与质量保障", [
        "单元测试", "组件测试", "API测试", "端到端测试", "支付回调测试",
        "并发测试", "安全测试", "兼容性测试", "移动真机测试", "回归测试",
        "用户验收测试", "缺陷分级",
    ]),
    "OBS": ("监控、运维与应急", [
        "可用性监控", "接口延迟监控", "错误率监控", "前端错误监控", "日志检索",
        "链路追踪", "告警分级", "值班制度", "故障通报", "根因分析",
        "容量评估", "运维月报",
    ]),
    "PER": ("性能、兼容与可访问性", [
        "首屏性能", "核心Web指标", "图片优化", "代码分包", "缓存策略",
        "弱网体验", "主流浏览器兼容", "响应式布局", "键盘可操作", "颜色对比度",
        "屏幕阅读器标签", "大字体适配",
    ]),
    "BCP": ("备份、容灾与业务连续性", [
        "数据库自动备份", "备份加密", "备份恢复演练", "跨可用区", "RPO目标",
        "RTO目标", "故障降级", "只读模式", "供应商接口熔断", "支付通道降级",
        "灾难恢复预案", "年度演练",
    ]),
    "PM": ("项目管理、培训与交付", [
        "需求基线", "原型评审", "UI评审", "技术方案评审", "迭代计划",
        "变更控制", "周报", "风险台账", "上线清单", "管理员培训",
        "客服培训", "验收与质保",
    ]),
}

API_GROUPS = {
    "身份与组织": ["POST /auth/sms/send", "POST /auth/login", "POST /auth/refresh", "GET /me", "GET /malls", "POST /malls/switch", "POST /sso/callback"],
    "商品与搜索": ["GET /products", "GET /products/{id}", "GET /categories", "GET /search/suggestions", "GET /search/hotwords", "POST /favorites", "DELETE /favorites/{id}"],
    "购物车与结算": ["GET /cart", "POST /cart/items", "PATCH /cart/items/{id}", "DELETE /cart/items/{id}", "POST /checkout/preview", "POST /orders", "POST /payments"],
    "订单与售后": ["GET /orders", "GET /orders/{id}", "POST /orders/{id}/cancel", "POST /orders/{id}/confirm", "POST /after-sales", "GET /after-sales", "POST /refunds"],
    "账户与卡券": ["GET /accounts", "GET /accounts/logs", "POST /accounts/grants", "GET /coupons", "POST /coupons/{id}/verify", "POST /coupon-provider/callback"],
    "管理与集成": ["POST /admin/products/import", "POST /admin/budgets/grant", "GET /admin/reports", "POST /suppliers/orders/push", "POST /logistics/callback", "POST /payments/callback", "GET /health"],
}

DATA_ENTITIES = [
    "platform", "tenant_group", "enterprise", "mall", "department", "employee", "user_identity",
    "role", "permission", "user_role", "supplier", "distributor", "brand", "category", "spu", "sku",
    "price_rule", "inventory", "promotion", "coupon_template", "user_coupon", "welfare_account",
    "account_budget_batch", "account_ledger", "cart", "cart_item", "address", "parent_order",
    "sub_order", "order_item", "payment", "refund", "settlement", "shipment", "verification_code",
    "after_sale", "invoice", "customer_ticket", "ticket_message", "operation_audit", "login_audit",
    "notification", "content_slot", "integration_job", "webhook_event",
]


def set_run_font(run, size=None, bold=None, color=None, east_asia="Microsoft YaHei"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.font.color.rgb = RGBColor.from_string(INK)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, NAVY, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run("智慧翼企业福利商城｜全量技术方案与报价评估报告")
    set_run_font(run, 8.5, False, MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("雍彻科技（YONGCHE TECH）  ·  内部评审/甲方沟通稿  ·  ")
    set_run_font(r, 8, False, MUTED)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r._r.append(fld_begin)
    r._r.append(instr)
    r._r.append(fld_end)


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(90)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("项目技术与商务评估报告")
    set_run_font(r, 12, True, ORANGE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("智慧翼企业福利商城")
    set_run_font(r, 29, True, NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run("全量技术方案、交付边界、客服制度与报价评估")
    set_run_font(r, 15, False, BLUE)

    for label, value in [
        ("拟用正式域名", "hbbtzn.com / www.hbbtzn.com（管理后台 smart.hbbtzn.com）"),
        ("商城展示名称", "智慧翼企业福利商城"),
        ("技术服务方", "雍彻科技（YONGCHE TECH）"),
        ("当前版本", "生产演示版 V2，Git 2d09a6c5，公开站点 HTTP 200"),
        ("报告日期", "2026年7月24日"),
        ("文档性质", "需求澄清、技术评审、报价与合同附件基础稿"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        lr = p.add_run(f"{label}：")
        set_run_font(lr, 10.5, True, MUTED)
        vr = p.add_run(value)
        set_run_font(vr, 10.5, False, INK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(55)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("重要说明：当前线上版本是可公开访问的前端演示原型，不是可承载真实资金与订单的正式生产商城。")
    set_run_font(r, 10.5, True, RED)
    doc.add_page_break()


def add_callout(doc, title, body, fill=LIGHT_BLUE, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    shade_cell(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{title}｜")
    set_run_font(r, 10.5, True, color)
    r = p.add_run(body)
    set_run_font(r, 10.5, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        shade_cell(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(text))
        set_run_font(r, font_size, True, NAVY)
    for row in rows:
        added_row = table.add_row()
        prevent_row_split(added_row)
        cells = added_row.cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            if idx in (0, 2, 3) and len(headers) >= 4:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(value))
            set_run_font(r, font_size, False, INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullets(doc, items: Iterable[str]):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{abstract_id + 1000:08X}")
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    level.append(lvl_text)
    p_pr = OxmlElement("w:pPr")
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "280")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_node)
        p._p.get_or_add_pPr().append(num_pr)
        r = p.add_run(item)
        set_run_font(r, 11, False, INK)


def add_numbered(doc, items: Iterable[str]):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"{abstract_id + 1000:08X}")
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "280")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)

    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.194)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_node)
        p._p.get_or_add_pPr().append(num_pr)
        r = p.add_run(item)
        set_run_font(r, 11, False, INK)


def add_para(doc, text, bold=False, color=INK):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, 11, bold, color)
    return p


def requirement_rows():
    all_rows = []
    for domain_index, (code, (domain, topics)) in enumerate(DOMAIN_TOPICS.items(), start=1):
        for index, topic in enumerate(topics, start=1):
            if index <= 6:
                priority, phase = "P0", "MVP"
            elif index <= 10:
                priority, phase = "P1", "正式V1"
            else:
                priority, phase = "P2", "二期"
            req_id = f"{code}-{index:03d}"
            requirement = f"系统必须支持{topic}，并在平台/集团/商城/员工四级上下文中执行权限与数据范围校验。"
            acceptance = f"以至少2个集团、2个商城、3类角色完成{topic}正向、越权、异常和审计验证；关键结果可追溯且数据不串租户。"
            all_rows.append((req_id, domain, topic, priority, phase, requirement, acceptance))
    return all_rows


def build_markdown(rows):
    lines = [
        "# 智慧翼企业福利商城全量技术方案与报价评估报告",
        "",
        "- 域名：hbbtzn.com / www.hbbtzn.com；管理后台 smart.hbbtzn.com",
        "- 商城名称：智慧翼企业福利商城",
        "- 技术服务方：雍彻科技（YONGCHE TECH）",
        "- 当前版本：V2，Git 2d09a6c5d61dc842d33b9ad84ceee98d5e82e5d3",
        "- 报告日期：2026-07-24",
        "",
        "## 结论",
        "",
        "当前版本是完成度较高、可公开访问的前端演示原型。路由、严格类型、商城级浏览器数据隔离、组合扣款上限、订单状态统一、售后状态写入、移动端底部导航、测试和依赖漏洞已修复；但真实上线仍需建设后端、数据库、管理端、认证、支付、供应商、物流、卡券、客服工单、审计与合规体系。",
        "",
        "## 推荐报价",
        "",
        "- 受控生产型MVP：建议对外含税报价52.8万元；目标成交48.8-52.8万元；低于42.8万元必须缩减范围。",
        "- 标准正式V1：98-138万元，周期3-4个月。",
        "- 完整多租户平台：180-320万元，周期6-9个月。",
        "- 一个月口径：4周开发冲刺；合同建议写6周（含30%缓冲与验收），且仅承诺P0范围。",
        "",
        f"## 全量需求清单（{len(rows)}项）",
        "",
        "|编号|领域|技术点|优先级|阶段|要求|验收|",
        "|---|---|---|---|---|---|---|",
    ]
    for row in rows:
        lines.append("|" + "|".join(str(value).replace("|", "｜") for value in row) + "|")
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def build_docx(rows):
    doc = Document()
    configure_document(doc)
    add_cover(doc)

    doc.add_heading("文档导航与使用方式", level=1)
    add_para(doc, "本报告可直接用于与甲方沟通、内部技术评审、估价以及后续拆分成需求规格说明书（SRS）、概要设计、接口规范、测试方案和运维手册。签约前必须将P0/P1/P2边界、第三方接口责任、数据迁移、验收口径和变更流程写入合同附件。")
    toc_rows = [
        ("1", "执行摘要与重新评分", "回答现在能否交付、修复了什么"),
        ("2", "当前代码审计与剩余边界", "区分演示版与正式生产系统"),
        ("3", "域名、名称与品牌治理", "域名、DNS、证书、备案与运营主体"),
        ("4", "业务与技术总体架构", "四级主体、三类业务管理层和系统分层"),
        ("5", "功能、数据、接口与安全设计", "核心业务闭环和后台能力"),
        ("6", "客服制度与运营SOP", "服务渠道、SLA、升级、质检和人员"),
        ("7", "项目计划、团队与报价", "一个月MVP口径、30%缓冲与费用"),
        ("8", "测试、验收、上线与运维", "验收标准和正式上线门槛"),
        ("附录A", f"全量技术需求矩阵（{len(rows)}项）", "可拆成甲方确认清单"),
        ("附录B", "接口目录与数据实体", "后端设计输入"),
        ("附录C", "风险、假设与合同保护条款", "避免无限责任和漏项"),
    ]
    add_table(doc, ["章节", "内容", "用途"], toc_rows, [950, 3600, 4810], 9)

    doc.add_heading("1. 执行摘要与重新评分", level=1)
    add_callout(doc, "总体结论", "修复后项目已从“有明显业务漏洞的前端演示稿”提升为“工程基础合格、可公开演示、可继续扩展的高保真原型”。它仍不能承载真实福利余额、真实支付、真实订单和真实卡券。")
    score_rows = [
        ("视觉与页面完整度", "7.0", "8.0", "蓝橙品牌统一，12个业务页面，新增移动端主导航；仍缺完整小程序设计系统"),
        ("前端工程质量", "4.0", "7.5", "严格TS、单元测试、状态统一、商城隔离、依赖0漏洞、可重复构建"),
        ("演示交付准备度", "—", "8.5", "生产公开地址HTTP 200，演示环境标识清晰"),
        ("正式上线准备度", "2.0", "3.5", "前端基础改善，但后端、资金、权限、管理端和集成仍未建设"),
        ("综合原型评分", "约4.3", "6.9", "适合作为正式项目的需求与前端基础，不等于生产商城"),
    ]
    add_table(doc, ["维度", "修复前", "修复后", "说明"], score_rows, [1700, 900, 900, 5860], 8.8)

    doc.add_heading("1.1 本次实际完成的修复", level=2)
    fixed = [
        "导航字段统一为currentPage，12个主页面可按Context路由切换。",
        "TypeScript strict=true，补齐React类型；当前tsc --noEmit无错误。",
        "福利卡、餐卡和外部补差统一按分计算，任意组合不得超过订单金额。",
        "多供应商拆单采用末单吸收尾差，父订单金额与子订单合计守恒。",
        "localStorage键增加mallId命名空间；购物车、订单、余额、流水、卡券、售后、地址、收藏按商城隔离。",
        "商城切换时同步刷新用户、购物车、订单、地址、收藏并返回首页，防止旧商城页面残留。",
        "首页、个人中心、订单页统一使用Context订单数据；下单、确认收货、售后后同步刷新。",
        "售后页真正调用服务层写入售后记录与订单状态，不再只显示成功文案。",
        "移除虚构ICP备案、400电话、Session与企业专线，明确标注演示环境。",
        "新增移动端五栏底部导航、移动头部、购物车角标、安全区域适配，并隐藏手机首屏的大型PC分类菜单。",
        "新增6个自动化测试，覆盖支付边界、金额精度、商城隔离、超额扣款和拆单守恒。",
        "升级构建依赖并完成npm audit，生产与开发依赖均为0个已知漏洞。",
    ]
    add_bullets(doc, fixed)

    doc.add_heading("1.2 验证证据", level=2)
    verify_rows = [
        ("Git提交", "2d09a6c5d61dc842d33b9ad84ceee98d5e82e5d3", "已推送托管源main"),
        ("生产版本", "Sites V2", "部署成功"),
        ("公开地址", "https://hbbtzn.com", "HTTP 200"),
        ("类型检查", "npm run lint", "通过"),
        ("自动测试", "2个测试文件、6个测试", "全部通过"),
        ("生产构建", "vinext build / Vite 8.1.5", "通过"),
        ("依赖审计", "npm audit --audit-level=low", "0漏洞"),
    ]
    add_table(doc, ["检查项", "结果", "状态"], verify_rows, [1800, 4560, 3000], 8.8)

    doc.add_heading("2. 当前代码审计与剩余边界", level=1)
    add_para(doc, "以下问题不是继续修改前端页面就能解决，而是正式项目必须新增的系统能力。报价、工期和合同验收必须以此为边界。")
    gap_rows = [
        ("真实数据可信性", "余额、订单、卡券仍存浏览器，用户可通过开发者工具修改", "建设服务端API、数据库、账本和审计"),
        ("身份与权限", "当前没有登录、SSO、角色和服务端授权", "建设IAM、RBAC/ABAC和会话管理"),
        ("管理后台", "没有平台、集团、商城、供应商后台", "建设至少三类后台与供应商工作台"),
        ("真实支付", "微信补差只是演示计算", "接入商户号、回调验签、对账、退款"),
        ("供应链", "商品、库存、物流、卡券均为Mock", "接入供应商并设计重试、幂等、补偿"),
        ("合规", "运营主体、备案、隐私、协议尚未由甲方确认", "甲方提供主体资料并完成法务审查"),
        ("移动端", "已有响应式H5基础，但不是微信小程序工程", "确定H5/小程序/App路线后独立交付"),
        ("测试深度", "已有关键单元测试，缺E2E、压测、安全和真机矩阵", "按上线门槛补齐完整测试体系"),
    ]
    add_table(doc, ["领域", "当前状态", "正式建设要求"], gap_rows, [1800, 3600, 3960], 8.8)

    doc.add_heading("3. 域名、名称与品牌治理", level=1)
    domain_rows = [
        ("主域名", "hbbtzn.com", "商城前台与业务 API 同源，已上线"),
        ("www域名", "www.hbbtzn.com", "建议 301 跳转至主域名"),
        ("管理后台", "smart.hbbtzn.com", "运营后台，独立部署"),
        ("商城展示名", "智慧翼企业福利商城", "建议甲方书面确认商标/名称使用权"),
        ("技术服务方", "雍彻科技（YONGCHE TECH）", "可展示于技术支持与版权说明"),
        ("运营主体", "待甲方确认", "必须决定备案、协议、收款、发票与客服责任主体"),
    ]
    add_table(doc, ["项目", "建议/现状", "治理要求"], domain_rows, [1600, 3300, 4460], 8.8)

    doc.add_heading("3.1 当前DNS待办", level=2)
    add_para(doc, "截至2026年7月24日最终复核，Sites侧证书状态已为active，但自定义域名仍处于pending：主域名访问返回500，www访问返回404，尚未稳定路由到本项目。需要在Cloudflare删除冲突的@与www旧A/AAAA记录，先以“仅DNS”方式添加以下记录；验证完成并确认主域名与www均返回200后，再决定是否启用代理。")
    dns_rows = [
        ("A", "@", "162.159.143.30"),
        ("A", "@", "172.66.3.26"),
        ("CNAME", "www", "custom-domains.chatgpt.site"),
        ("TXT", "_openai-site-verification", "openai-site-verification=M_9k3OeUkf5hKCrUtJeL4cbrRqr6LgcoX3lZtTQKgqE"),
        ("TXT", "_cf-custom-hostname", "9337d7ca-44b3-49ac-8a72-c6201b0f62c5"),
        ("TXT", "_openai-site-verification.www", "openai-site-verification=WajUbgv-2PcGvmq0e7sNYrzc9jx2Giftbf5THFjEcbA"),
        ("TXT", "_cf-custom-hostname.www", "bfde253a-f891-4bf5-8c17-03bde8ec8f19"),
    ]
    add_table(doc, ["类型", "名称", "值"], dns_rows, [900, 2700, 5760], 8)

    doc.add_heading("4. 业务与技术总体架构", level=1)
    add_callout(doc, "四级业务主体", "平台 → 集团 → 商城 → 员工。管理责任通常体现为平台运营、集团管理、商城运营三类业务管理层；员工是消费与服务对象，不算管理层。因此可以表述为“四级业务主体、三级管理层”。")
    add_para(doc, "建议的正式系统分为体验层、业务服务层、平台能力层和基础设施层。所有服务端查询必须自动注入tenant_id、enterprise_id和mall_id，不能依赖前端传参实现隔离。")
    architecture_rows = [
        ("体验层", "PC Web、移动H5、微信小程序、平台后台、集团后台、商城后台、供应商工作台"),
        ("业务服务层", "用户、商品、搜索、价格、促销、账户、购物车、订单、支付、履约、卡券、售后、客服"),
        ("平台能力层", "IAM、租户、审计、消息、任务、文件、配置、API网关、报表、风控"),
        ("数据层", "关系数据库、Redis、对象存储、搜索引擎、消息队列、日志与指标"),
        ("基础设施层", "DNS/CDN/WAF、容器或Serverless、CI/CD、密钥、备份、监控、容灾"),
    ]
    add_table(doc, ["层级", "组成"], architecture_rows, [1700, 7660], 9)

    doc.add_heading("4.1 关键架构原则", level=2)
    add_bullets(doc, [
        "账本优先：余额不得只保存一个可修改数字，必须使用不可变流水和事务生成余额快照。",
        "服务端授权：任何订单、地址、卡券、余额查询都必须在服务端校验用户与商城归属。",
        "幂等优先：下单、扣款、支付回调、退款、发码和供应商推单均需要幂等键。",
        "事件驱动：订单状态、物流、发码、退款和通知用事件及重试队列解耦。",
        "可观测：每个请求携带trace_id，资金和订单操作必须能从用户追溯到外部回调。",
        "配置化：集团、商城、支付方式、账户适用范围、客服SLA和供应商路由均应后台配置。",
    ])

    doc.add_heading("5. 功能、数据、接口与安全设计", level=1)
    doc.add_heading("5.1 核心交易闭环", level=2)
    add_numbered(doc, [
        "员工通过手机号、微信或企业SSO登录；服务端确定其集团、商城、角色和账户。",
        "商品中心按商城、供应商、区域和账户适用范围返回可见商品与实时价格库存。",
        "购物车保存SKU快照，但结算时必须重新校验价格、库存、限购、地址和促销。",
        "结算服务按供应商和履约方式拆单，账户服务预占福利资金，支付服务处理外部补差。",
        "支付成功后订单进入履约；虚拟商品发码，实物商品推送供应商并订阅物流。",
        "售后按子订单和支付构成计算退款，福利与餐卡原路退回对应账本，微信原路退款。",
        "每天执行支付、账户、订单、供应商结算四方对账，差异进入财务工单。",
    ])

    doc.add_heading("5.2 数据模型", level=2)
    add_para(doc, f"正式数据库至少包含{len(DATA_ENTITIES)}个核心实体。每个业务表建议包含id、tenant_id、enterprise_id、mall_id、created_at、updated_at、created_by、version和软删除/状态字段；资金流水不允许物理删除。")
    entity_rows = []
    for index in range(0, len(DATA_ENTITIES), 4):
        chunk = DATA_ENTITIES[index:index + 4]
        entity_rows.append(tuple(chunk + [""] * (4 - len(chunk))))
    add_table(doc, ["实体1", "实体2", "实体3", "实体4"], entity_rows, [2340, 2340, 2340, 2340], 8.5)

    doc.add_heading("5.3 API设计", level=2)
    api_rows = []
    for group, endpoints in API_GROUPS.items():
        for endpoint in endpoints:
            api_rows.append((group, endpoint, "JWT/SSO + tenant context", "幂等、审计、统一错误码"))
    add_table(doc, ["分组", "示例接口", "认证", "横切要求"], api_rows, [1500, 3500, 2100, 2260], 8)

    doc.add_heading("5.4 安全基线", level=2)
    add_bullets(doc, [
        "敏感数据采用TLS传输；卡密、身份证件、手机号等按字段加密，密钥不进入代码仓库。",
        "福利账户扣减使用数据库事务、乐观锁/行锁和幂等凭证，禁止前端直接写余额。",
        "管理员采用MFA、最小权限、IP/设备策略和高风险操作二次确认。",
        "API网关执行签名、时间戳、nonce、限流、黑名单与回放防护。",
        "日志不得记录完整卡密、验证码、支付签名、身份证号和访问令牌。",
        "上线前至少完成依赖扫描、SAST、DAST、越权专项、支付回调专项和隐私合规检查。",
    ])

    doc.add_heading("6. 客服制度与运营SOP", level=1)
    add_callout(doc, "责任边界", "客服制度必须由甲方运营主体确认。雍彻科技可提供系统、工单、知识库、报表和技术二线支持；商品承诺、退款裁决、发票和资金责任应由运营主体及供应商承担。")
    service_rows = [
        ("服务渠道", "商城在线客服、企业微信、电话、邮件、工单", "所有渠道最终归集同一工单"),
        ("服务时间", "工作日09:00-18:00；大促/发薪日可延长", "非服务时间机器人收单并提示响应时间"),
        ("L1一线", "咨询、订单查询、规则解释、标准售后", "使用知识库，不得越权修改资金"),
        ("L2业务", "复杂售后、供应商催办、财务对账", "商城运营/供应链/财务处理"),
        ("L3技术", "系统故障、数据异常、支付回调、批量事故", "雍彻科技技术支持"),
        ("重大事件", "资金错扣、批量卡密泄露、全站不可用", "15分钟升级，立即止损并通报"),
    ]
    add_table(doc, ["制度项", "建议", "控制要求"], service_rows, [1700, 3600, 4060], 8.8)

    doc.add_heading("6.1 工单优先级与SLA", level=2)
    sla_rows = [
        ("P0 紧急", "全站不可用、批量资金错误、数据泄露", "15分钟", "2小时止损/4小时恢复", "电话+群组+管理层"),
        ("P1 高", "支付失败、批量订单/发码异常", "30分钟", "4小时临时方案/8小时恢复", "L2+L3"),
        ("P2 中", "单用户订单、退款、物流异常", "2小时", "1个工作日", "L1→L2"),
        ("P3 低", "咨询、建议、非阻断缺陷", "4小时", "3个工作日", "L1"),
    ]
    add_table(doc, ["级别", "场景", "首次响应", "目标", "升级"], sla_rows, [950, 2700, 1300, 2300, 2110], 8.5)

    doc.add_heading("6.2 标准客服流程", level=2)
    add_numbered(doc, [
        "受理：自动生成工单号，记录集团、商城、员工、订单、渠道与用户诉求。",
        "身份核验：只询问必要信息；涉及余额、卡密和地址变更必须二次核验。",
        "分类定级：按咨询、订单、物流、卡券、支付、退款、发票、投诉和系统故障分类。",
        "首响与告知：说明当前状态、预计处理时间和下一次反馈时间，不作无授权承诺。",
        "处理与协同：调用供应商、财务或技术子工单；主工单对用户保持单一窗口。",
        "结果确认：提供处理结果、退款去向、到账时间或替代方案，并让用户确认。",
        "关闭与评价：用户确认或超时自动关闭，收集满意度；低分自动复核。",
        "复盘：P0/P1及重复问题进入知识库、供应商考核和产品改进清单。",
    ])

    doc.add_heading("6.3 客服质检与KPI", level=2)
    kpi_rows = [
        ("首次响应达标率", "≥95%", "按优先级SLA统计"),
        ("一次解决率", "≥75%", "无二次转派且用户确认"),
        ("工单按时解决率", "≥92%", "扣除用户/供应商等待需留证"),
        ("用户满意度", "≥4.5/5", "样本量和评价率同时披露"),
        ("重复咨询率", "逐月下降", "驱动页面、规则和知识库优化"),
        ("资金类差错", "0容忍", "一旦发生按P0处理"),
    ]
    add_table(doc, ["指标", "建议目标", "口径"], kpi_rows, [2200, 1700, 5460], 8.8)

    doc.add_heading("7. 项目计划、团队与报价", level=1)
    add_callout(doc, "一个月MVP的正确口径", "可以做4周开发冲刺，但合同工期建议写6周：4周实施 + 30%缓冲（约1.2周）+ 验收/上线窗口。若要求自然月内正式上线，必须冻结P0范围、接口按时提供，并配备8-10人并行团队。", "FFF4E5", ORANGE)
    phase_rows = [
        ("第0周", "需求冻结、接口清单、原型与验收口径", "甲方签字确认P0和排除项"),
        ("第1周", "后端底座、租户/IAM、数据库、前后台框架", "登录、四级隔离、基础后台"),
        ("第2周", "商品、购物车、账户、订单、管理端", "主交易流程联调"),
        ("第3周", "支付/供应商/物流或卡券最小接口、售后", "端到端闭环"),
        ("第4周", "回归、安全、数据初始化、培训、预发布", "候选上线版本"),
        ("缓冲1.2周", "接口延期、缺陷、甲方反馈", "不新增范围"),
        ("验收0.8周", "UAT、上线审批、DNS/证书、观察期", "签署验收单"),
    ]
    add_table(doc, ["时间", "工作", "出口"], phase_rows, [1500, 4460, 3400], 8.8)

    doc.add_heading("7.1 MVP团队与人月测算", level=2)
    team_rows = [
        ("项目/产品经理", "1", "1.2", "需求、范围、计划、验收、甲方协调"),
        ("UI/UX", "1", "0.8", "PC+移动关键页面与设计系统"),
        ("前端工程师", "2", "2.0", "商城前台与管理后台"),
        ("后端工程师", "3", "3.0", "租户、账户、订单、接口、后台"),
        ("测试工程师", "1-2", "1.3", "功能、接口、回归、上线验收"),
        ("DevOps/安全", "0.5-1", "0.6", "环境、CI/CD、监控、安全基线"),
        ("合计", "8-10", "8.9人月", "一个月并行投入"),
    ]
    add_table(doc, ["角色", "人数", "人月", "职责"], team_rows, [2100, 1000, 1300, 4960], 8.8)

    doc.add_heading("7.2 建议报价", level=2)
    price_rows = [
        ("现有演示原型及工程修复", "3-6万元", "仅前端演示、部署、文档，不含真实业务后台"),
        ("受控生产型MVP（P0）", "建议52.8万元含税", "目标成交48.8-52.8万元；底价42.8万元且必须减范围"),
        ("标准正式V1（P0+主要P1）", "98-138万元", "PC+移动+多后台+核心集成，3-4个月"),
        ("完整多租户福利平台", "180-320万元", "多供应商、结算、BI、风控、容灾，6-9个月"),
        ("年度维保", "软件合同额15%-20%/年", "缺陷、升级、监控、应急；新需求另计"),
        ("客服运营人力", "2.5-4.5万元/月", "2名一线+兼职主管参考，不含大促扩容"),
    ]
    add_table(doc, ["方案", "建议价格", "边界"], price_rows, [2500, 2400, 4460], 8.8)
    add_para(doc, "报价策略建议：先报52.8万元的受控生产型MVP，并将第三方费用、真实支付/供应商资质、甲方接口改造、历史数据清洗、微信认证与备案、短信和客服人力列为甲方承担或单独报价。不要以当前前端页面数量直接估价；正式成本主要在后端资金安全、数据隔离、管理后台、接口联调、测试和验收。", True, NAVY)

    doc.add_heading("7.3 MVP费用拆分", level=2)
    cost_rows = [
        ("需求/产品/项目管理", "6.0万元", "范围、原型、验收、会议、风险"),
        ("UI/UX与移动适配", "3.5万元", "设计系统和关键移动流程"),
        ("前端商城与后台", "9.0万元", "PC/H5与管理端"),
        ("后端与数据库", "14.0万元", "租户、账户、订单、后台、接口"),
        ("第三方集成", "6.0万元", "限定2类核心接口，超出另计"),
        ("测试与安全", "5.0万元", "自动化、回归、安全、UAT"),
        ("DevOps与上线", "2.5万元", "环境、部署、监控、DNS"),
        ("30%风险与交付缓冲", "6.8万元", "接口不确定性、验收、返工"),
        ("合计", "52.8万元", "建议含税对外总价"),
    ]
    add_table(doc, ["成本项", "金额", "说明"], cost_rows, [2800, 1800, 4760], 8.8)

    doc.add_heading("7.4 第三方与持续费用（不建议包死）", level=2)
    add_bullets(doc, [
        "域名、备案、商标、微信认证、支付商户、短信签名等主体费用由甲方承担。",
        "云资源、CDN、数据库、对象存储、日志、WAF按流量计费；初期可预估1.2-6万元/年。",
        "短信、实名认证、地图、物流查询、电子发票、卡券供应商按调用量或交易额计费。",
        "支付通道手续费、供应商结算、退款手续费和资金监管不属于软件开发费。",
        "新增供应商接口建议按1.5-5万元/家评估；非标准或缺文档接口单独估时。",
        "微信小程序认证、上架整改、隐私合规整改应保留1-3轮，超出约定轮次走变更。",
    ])

    doc.add_heading("7.5 商务与付款建议", level=2)
    add_table(doc, ["节点", "比例", "条件"], [
        ("合同签署/启动", "30%", "人员排期和环境准备"),
        ("原型与技术方案确认", "30%", "P0范围、接口、验收基线签字"),
        ("预发布/UAT", "30%", "候选版本部署并进入验收"),
        ("终验与移交", "10%", "验收单、源码、文档、培训完成"),
    ], [2500, 1400, 5460], 8.8)

    doc.add_heading("8. 测试、验收、上线与运维", level=1)
    doc.add_heading("8.1 正式上线门槛", level=2)
    add_bullets(doc, [
        "P0需求全部有测试用例、负责人和验收结果，P1/P2未完成项不得口头混入验收。",
        "至少完成两个集团、两个商城、多个角色的越权与数据串租户专项测试。",
        "资金账本、支付回调、重复下单、退款、拆单尾差与对账通过专项测试。",
        "关键接口具备超时、重试、幂等、熔断和人工补偿路径。",
        "完成浏览器、移动真机、弱网、性能、安全、备份恢复和故障演练。",
        "隐私政策、用户协议、客服渠道、发票、备案和运营主体信息已确认。",
        "生产监控、告警、值班、回滚、应急联系人和变更窗口已就绪。",
    ])
    doc.add_heading("8.2 建议验收指标", level=2)
    acceptance_rows = [
        ("可用性", "月可用性≥99.9%（计划维护除外）"),
        ("性能", "核心页面LCP≤2.5秒目标；核心API P95≤500ms（不含第三方）"),
        ("错误率", "核心API 5xx低于0.1%"),
        ("安全", "无未处置的严重/高危漏洞；越权与资金专项通过"),
        ("隔离", "跨集团/商城数据访问测试100%阻断"),
        ("账务", "订单、账户、支付对账差异可解释且闭环"),
        ("测试", "P0用例通过率100%，P1≥98%，无阻断缺陷"),
        ("文档", "部署、接口、数据字典、运维、客服与培训资料齐全"),
    ]
    add_table(doc, ["维度", "建议门槛"], acceptance_rows, [2200, 7160], 9)

    doc.add_heading("附录A. 全量技术需求矩阵", level=1)
    add_callout(doc, "使用方法", f"本矩阵共{len(rows)}项。签约前由甲方逐项勾选MVP/正式V1/二期或不做，并补充责任人、数据来源和验收证据。当前优先级是雍彻科技基于通用福利商城给出的建议，不代替甲方书面确认。")
    for code, (domain, topics) in DOMAIN_TOPICS.items():
        doc.add_heading(f"{code}｜{domain}", level=2)
        domain_rows = []
        for row in rows:
            if row[0].startswith(code + "-"):
                domain_rows.append((row[0], row[2], row[3], row[4], row[5], row[6]))
        add_table(
            doc,
            ["编号", "技术点", "优先级", "阶段", "要求", "验收标准"],
            domain_rows,
            [950, 1500, 700, 850, 2560, 2800],
            7.2,
        )

    doc.add_heading("附录B. 接口、数据与非功能清单", level=1)
    doc.add_heading("B.1 接口统一规则", level=2)
    add_bullets(doc, [
        "所有写接口接受Idempotency-Key；返回统一code、message、data、trace_id。",
        "分页统一page、page_size或cursor，禁止不同接口自行发明分页格式。",
        "外部回调必须验签、校验时间戳与nonce，原始报文脱敏留档。",
        "接口版本通过URL或Header管理；废弃接口至少提前一个发布周期通知。",
        "批量任务异步化，返回job_id并提供进度、结果与失败明细下载。",
        "OpenAPI文档、示例、错误码、限流、重试和联系人必须随版本更新。",
    ])
    doc.add_heading("B.2 数据保留建议", level=2)
    add_table(doc, ["数据", "建议保留", "说明"], [
        ("订单、支付、退款、结算", "按财税和合同要求，建议不少于5年", "最终由甲方法务/财务确定"),
        ("账户账本与资金审计", "长期或依法保留", "只追加、不可覆盖"),
        ("管理员审计", "不少于1年", "高风险操作可延长"),
        ("登录与安全日志", "不少于6个月", "结合等保和安全要求"),
        ("客服工单", "2-3年", "涉及争议的延长至结案后"),
        ("临时验证码/令牌", "分钟至小时级", "到期自动删除"),
    ], [2600, 2600, 4160], 8.8)

    doc.add_heading("附录C. 风险、假设与合同保护条款", level=1)
    risk_rows = [
        ("需求持续变化", "高", "P0基线签字；超范围走CR并调整费用工期"),
        ("甲方接口延期", "高", "接口占位Mock；延期顺延，不承担窝工"),
        ("运营主体不明确", "高", "备案、协议、支付、发票、客服主体在启动前确认"),
        ("供应商接口质量", "高", "设置重试、补偿和人工工单；供应商SLA由甲方推动"),
        ("真实资金风险", "高", "账本、事务、对账、审计和灰度；不得沿用localStorage"),
        ("一个月工期", "高", "限定P0、8-10人并行、甲方每日决策；合同写6周"),
        ("历史数据迁移", "中", "样例数据和字段映射先验收；清洗另计"),
        ("小程序审核", "中", "平台审核时间不计入开发可控工期"),
        ("域名与备案", "中", "甲方提供Cloudflare/备案权限和主体资料"),
        ("验收拖延", "中", "约定验收窗口与逾期视为阶段验收"),
    ]
    add_table(doc, ["风险", "级别", "合同/技术措施"], risk_rows, [2200, 900, 6260], 8.8)
    doc.add_heading("C.1 必须写入合同的排除项", level=2)
    add_bullets(doc, [
        "未经书面列出的第三方系统、供应商、历史数据源和硬件设备。",
        "甲方或第三方账号、资质、备案、商户号、短信签名、商标及内容合法性。",
        "第三方平台审核时长、接口故障、政策变化及收费调整。",
        "7×24客服人员、商品运营、供应商履约和财务结算执行，除非另签运营服务。",
        "超出约定容量、并发、数据量和服务区域后的扩容与架构升级。",
        "因甲方自行改代码、改DNS、泄露密钥或绕过流程造成的问题。",
    ])

    doc.add_heading("最终建议", level=1)
    add_callout(doc, "建议决策", "将现有V2作为甲方需求确认与演示基线；先签“受控生产型MVP”而不是承诺完整商城。建议对外报价52.8万元，合同工期6周，开发冲刺4周，预留30%缓冲；任何真实资金、支付、供应商和运营责任均按本报告边界落入合同。", LIGHT_BLUE, NAVY)
    add_para(doc, "报告编制：雍彻科技（YONGCHE TECH）", True, NAVY)
    add_para(doc, "版本：1.0｜日期：2026年7月24日｜本报告为技术与商务评估建议，最终范围、价格、税务和法律责任以双方盖章合同及附件为准。", False, MUTED)

    doc.core_properties.title = "智慧翼企业福利商城全量技术方案与报价评估报告"
    doc.core_properties.subject = "域名、需求、架构、客服制度、项目计划、报价与验收"
    doc.core_properties.author = "雍彻科技（YONGCHE TECH）"
    doc.core_properties.keywords = "智慧翼, 福利商城, hbbtzn.com, 雍彻科技, MVP, 技术方案, 报价"
    doc.save(DOCX_PATH)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    rows = requirement_rows()
    build_markdown(rows)
    build_docx(rows)
    print(f"requirements={len(rows)}")
    print(f"docx_bytes={DOCX_PATH.stat().st_size}")
    print(f"markdown_bytes={MD_PATH.stat().st_size}")


if __name__ == "__main__":
    main()
